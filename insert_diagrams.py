"""
Insert 12 diagram images into the training report at appropriate positions.
Reads from original document, writes to a separate output file.
"""
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOC_PATH = r"C:\Users\24522\Desktop\AI网络安全管家_实训报告_陈梓烨.docx"
IMG_DIR = r"C:\Users\24522\Desktop\图"
OUTPUT_PATH = r"C:\Users\24522\Desktop\AI网络安全管家_实训报告_陈梓烨_含图.docx"

DIAGRAMS = [
    # === Chapter 4 - 系统概要设计 ===
    {"file": "01_系统整体架构图.drawio.png",         "caption": "图4-1  系统整体架构图",           "search": "图4-1  系统整体架构图",          "replace": True},
    {"file": "02_系统主流业务流程图.drawio.png",       "caption": "图4-2  系统主流业务流程图",        "search": "图4-2  系统主流业务流程图",        "replace": True},
    {"file": "03_系统功能模块图.drawio.png",          "caption": "图4-3  系统功能模块图",           "search": "4.3  系统业务流程设计",           "replace": False},
    {"file": "12_知识库ER图.drawio.png",             "caption": "图4-4  知识库ER图",              "search": "5  系统详细设计",               "replace": False},
    # === Chapter 5 - 系统详细设计 ===
    {"file": "07_系统数据流图_DFD.drawio.png",        "caption": "图5-1  系统数据流图（DFD）",        "search": "5.1.2  NetworkMonitor",       "replace": False},
    {"file": "08_系统部署图.drawio.png",              "caption": "图5-2  系统部署图",               "search": "5.1.2  NetworkMonitor",       "replace": False},
    {"file": "10_防火墙多级降级策略流程图.drawio.png",   "caption": "图5-3  防火墙多级降级策略流程图",     "search": "5.2  AI Agent",              "replace": False},
    {"file": "04_ReAct核心循环流程图.drawio.png",      "caption": "图5-4  ReAct核心循环流程图",        "search": "5.2.2  工具注册表",             "replace": False},
    {"file": "09_MCP协议交互时序图.drawio.png",        "caption": "图5-5  MCP协议交互时序图",          "search": "5.2.3  安全沙箱",              "replace": False},
    {"file": "05_安全沙箱6层防护架构图.drawio.png",     "caption": "图5-6  安全沙箱6层防护架构图",       "search": "5.3  ESP32",                "replace": False},
    {"file": "11_知识库RAG流程图.drawio.png",          "caption": "图5-7  知识库RAG流程图",            "search": "5.3  ESP32",                "replace": False},
    {"file": "06_PC与ESP32通信时序图.drawio.png",      "caption": "图5-8  PC与ESP32通信时序图",        "search": "5.4  ",                     "replace": False},
]


def make_image_paragraph(doc, img_path):
    """Create a centered paragraph with an image (appended at end of doc)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    return para


def make_caption_paragraph(doc, caption_text):
    """Create a centered caption paragraph (appended at end of doc)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(caption_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.name = "宋体"
    rPr = run._element.get_or_add_rPr()
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)
    return para


def find_last_paragraph_by_text(doc, search_text):
    """
    Find the LAST paragraph containing search_text.
    Returns the paragraph index, or None if not found.
    Using the LAST match ensures we target the actual section heading,
    not the TOC entry (which appears earlier in the document).
    """
    matches = []
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            matches.append(i)
    return matches[-1] if matches else None


def main():
    # Verify input exists
    if not os.path.exists(DOC_PATH):
        print(f"ERROR: Document not found: {DOC_PATH}")
        print("Please restore the original document and try again.")
        return

    doc = docx.Document(DOC_PATH)
    print(f"Opened: {DOC_PATH}")
    print(f"Original paragraph count: {len(doc.paragraphs)}")

    # Validate all images
    for diag in DIAGRAMS:
        img_path = os.path.join(IMG_DIR, diag["file"])
        if not os.path.exists(img_path):
            print(f"WARNING: {diag['file']} NOT FOUND in {IMG_DIR}")

    # Phase 1: Find targets
    insertions = []
    for diag in DIAGRAMS:
        img_path = os.path.join(IMG_DIR, diag["file"])
        if not os.path.exists(img_path):
            continue

        idx = find_last_paragraph_by_text(doc, diag["search"])
        if idx is None:
            print(f"SKIP: search text not found: '{diag['search'][:50]}'")
            continue

        context = doc.paragraphs[idx].text[:60]
        insertions.append({"idx": idx, "diag": diag, "img_path": img_path})
        print(f"  P{idx:>4d}: '{context}'")
        print(f"         -> {diag['file']} [{diag['caption']}]")

    # Phase 2: Sort bottom-to-top
    insertions.sort(key=lambda x: x["idx"], reverse=True)

    # Phase 3: Execute insertions
    for ins in insertions:
        diag = ins["diag"]
        img_path = ins["img_path"]

        # Re-find target after previous insertions shifted indices
        current_idx = find_last_paragraph_by_text(doc, diag["search"])
        if current_idx is None:
            print(f"ERROR: Lost target: '{diag['search'][:50]}'")
            continue

        target_elem = doc.paragraphs[current_idx]._element

        # Create paragraphs at document end
        img_para = make_image_paragraph(doc, img_path)
        cap_para = make_caption_paragraph(doc, diag["caption"])

        if diag["replace"]:
            # Remove old placeholder caption
            parent = target_elem.getparent()
            next_sib = target_elem.getnext()
            parent.remove(target_elem)

            # Insert image then caption above next_sibling
            # (addprevious inserts RIGHT BEFORE the target element)
            if next_sib is not None:
                # Insert image first, then caption — so caption ends up AFTER image
                next_sib.addprevious(img_para._element)
                next_sib.addprevious(cap_para._element)
            else:
                parent.append(img_para._element)
                parent.append(cap_para._element)
        else:
            # Insert image then caption before target paragraph
            # IMPORTANT: image first, then caption
            # After both calls: ... img ... cap ... target (top-to-bottom)
            target_elem.addprevious(img_para._element)
            target_elem.addprevious(cap_para._element)

        print(f"  {'R' if diag['replace'] else 'I'} {diag['caption']}")

    # Save to separate output file (don't overwrite original)
    doc.save(OUTPUT_PATH)
    print(f"\n{'='*60}")
    print(f"Done! {len(insertions)} diagrams inserted.")
    print(f"Output: {OUTPUT_PATH}")
    print(f"(Original preserved at: {DOC_PATH})")


if __name__ == "__main__":
    main()
