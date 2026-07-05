#!/usr/bin/env python3
"""Generate polished resume DOCX"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.line_spacing = 1.15
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
rPr.insert(0, rFonts)

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): '1a56db',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_project_title(doc, name, role):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.font.size = Pt(11)
    run.font.bold = True
    run2 = p.add_run("  |  " + role)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    for run in p.runs:
        run.font.size = Pt(10)

def add_tags(doc, tags):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for tag in tags:
        run = p.add_run("  " + tag + "  ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

# ========== HEADER ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("陈 梓 烨")  # 陈梓烨
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
contact = "\U0001f4e7 2452262230@qq.com  |  \U0001f4f1 132-6803-0361  |  \U0001f4cd 广东省广州市  |  2005.04 / 本科在读"
run = p.add_run(contact)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ========== EDUCATION ==========
add_section_title(doc, "\U0001f4da 教育背景")

p = doc.add_paragraph()
run = p.add_run("新疆大学    软件工程    本科在读    ")
run.font.size = Pt(10.5)
run.font.bold = True
run2 = p.add_run("2023.09 - 2027.06")
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_bullet(doc, "主修课程: 数据结构与算法、操作系统、计算机网络、数据库原理、软件工程、C/Python 程序设计")
add_bullet(doc, "英语能力: CET-4，高考英语 120+，具备英文技术文档和 API 参考独立阅读能力")

# ========== SKILLS ==========
add_section_title(doc, "\U0001f6e0️ 技术能力")

skills = [
    ("编程语言", "Python (熟练，tkinter/Flask/psutil/pywin32/anthropic SDK 等生态)，C/C++ (数据结构与算法)"),
    ("LLM/AI 应用", "大模型应用架构设计，多 Provider 统一抽象层（Anthropic/DeepSeek/智谱/通义千问/Moonshot/硅基流动），Anthropic SDK 原生集成 + OpenAI 兼容协议适配，Tool-Use Agent 循环，Prompt Engineering"),
    ("嵌入式/IoT", "ESP32 微控制器开发 (Arduino C++)，I2C/SPI/UART 通信协议，OLED/TFT 显示驱动，PWM 外设控制"),
    ("桌面应用开发", "Tkinter + ttkbootstrap 现代主题 GUI 开发，多线程并发与 UI 分离，Windows COM 自动化"),
    ("Web 后端", "Flask RESTful API 设计，WebSocket 实时通信，Jinja2 模板渲染"),
    ("计算机视觉", "Tesseract OCR 中英文混合识别，Pillow/PyMuPDF 图像处理与文档布局分析"),
    ("系统安全", "Windows EventLog 安全审计，网络连接实时监控与威胁检测，进程行为分析，防火墙与 Defender 状态巡检"),
    ("开发工具链", "Git 版本控制，Arduino IDE，VS Code，pip/virtualenv 环境管理，Poetry/pyproject.toml 打包"),
]

for sn, sd in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(sn + ": ")
    run.font.size = Pt(10)
    run.font.bold = True
    run2 = p.add_run(sd)
    run2.font.size = Pt(10)

# ========== PROJECTS ==========
add_section_title(doc, "\U0001f4bc 项目经验")

# Project 1: Code Review Agent
add_project_title(doc, "Code Review Agent — 多模型 AI 代码审查工具", "独立开发 | 2025.06")
add_tags(doc, ["Python", "Anthropic SDK", "OpenAI Protocol", "Git", "Tool-Use Agent", "CLI", "Rich", "7 Providers"])

add_bullet(doc, "设计并实现专业级 AI 驱动代码审查工具，支持对 Git Diff 进行深度分析，自动识别 Bug 风险、安全漏洞、性能问题、代码风格和维护性问题五大类缺陷")
add_bullet(doc, "构建 LLM Provider 抽象层，基于策略模式实现 BaseLLMClient 统一接口，分别以 Anthropic 原生 SDK 和 OpenAI 兼容协议两种后端实现；预置 7 个国内外模型 Provider（Anthropic Claude / DeepSeek / 通义千问 / 智谱 GLM / Moonshot Kimi / 硅基流动 / 自定义兼容端点），支持通过环境变量或 CLI 参数灵活切换")
add_bullet(doc, "在 OpenAI 兼容客户端中实现了 Anthropic ↔ OpenAI 工具调用协议的双向转换（Tool Schema 格式映射 + Message Content Block 翻译 + Tool Result 路由），使得 DeepSeek 等非 Anthropic 模型也能在 Tool-Use Agent 循环中正确调用工具")
add_bullet(doc, "实现 Tool-Use Agent 循环引擎：LLM 在审查过程中可通过 read_file 获取完整文件上下文、search_pattern (ripgrep) 跨文件搜索关联代码、git_blame 追溯历史变更，最大迭代 10 轮直到模型完成审查")
add_bullet(doc, "自研 Git Diff 解析器，将 git diff 输出解析为 DiffFile/DiffHunk 结构化对象，支持新增/修改/删除/重命名四种文件状态识别；实现智能分块算法（500行/块），确保大体积 Diff 适配各模型 Context Window")
add_bullet(doc, "实现结构化输出解析（正则提取 FINDING/END_FINDING 标记块 → Finding 数据类），支持三种输出格式：Rich 终端彩色面板、JSON（CI/自动化流水线友好）、Markdown（文档归档）")
add_bullet(doc, "完整的 Python 项目工程化：pyproject.toml 打包配置、argparse CLI 子命令体系、pip install 一键安装、Claude Code 技能系统深度集成")

# Project 2: AI Security Guardian
add_project_title(doc, "AI 网络安全管家 — ESP32 嵌入式智能终端", "独立开发 | 2025.06")
add_tags(doc, ["ESP32", "Arduino C++", "Python", "Flask", "LLM", "psutil", "pywin32", "I2C OLED"])

add_bullet(doc, "设计并实现基于 ESP32 微控制器的桌面安全监控设备，搭载 0.96寸 OLED 屏通过 USB Serial 与 PC 实时通信，全部元器件可淘宝购买并手工组装（硬件成本仅 ¥33）")
add_bullet(doc, "PC 端 Python Agent 集成四大安全引擎：Windows EventLog 安全审计、实时网络连接监控（识别 20+ 恶意端口/C2/可疑 IP 段）、进程行为分析、防火墙与 Defender 状态巡检")
add_bullet(doc, "自研 40×48px 程序化像素角色动画引擎，实现二次元陪伴角色，支持 8 种表情 + 眨眼/弹跳/抖动物理动画 + 对话气泡系统")
add_bullet(doc, "构建统一多模型 LLM 客户端（DeepSeek/智谱/硅基流动），自动故障转移 + Token 统计，Flask Web Dashboard + 系统托盘常驻")

# Project 3: Document Converter
add_project_title(doc, "智能文档格式转换工具 (PDF/PPTX/DOCX ↔ Word/Markdown)", "独立开发 | 2025.05")
add_tags(doc, ["Python", "Tkinter", "ttkbootstrap", "PyMuPDF", "Tesseract OCR", "多线程"])

add_bullet(doc, "开发跨格式文档转换桌面应用，采用调度器模式实现模块解耦，8 个功能模块独立可复用")
add_bullet(doc, "集成 Tesseract OCR 引擎，自动检测多平台安装路径，中英文混合识别；文本量阈值算法智能判定纯图片页面，自动触发 OCR")
add_bullet(doc, "基于 ttkbootstrap 现代主题的完整 GUI，多线程后台转换避免 UI 阻塞，实时日志 + 进度反馈 + 自动定位输出文件")

# Campus Projects
add_project_title(doc, "电商平台交易数据分析系统", "项目组长 (4人团队) | 2024.06")
add_tags(doc, ["Python", "Pandas", "Matplotlib", "数据清洗", "团队管理"])

add_bullet(doc, "带领 4 人团队完成数据分析全流程，使用 Pandas 处理电商订单数据集，输出多维度统计分析和可视化图表，独立撰写六千字项目论文并通过答辩")

add_project_title(doc, "校园管理系统 & 倒计时应用 (中软国际实训)", "项目成员 | 2024.03 - 2024.05")
add_tags(doc, ["C语言", "数据结构", "企业实训"])

add_bullet(doc, "在中软国际指导下完成校园管理系统、电子日历、倒计时等项目的全流程开发，掌握 C 语言编程、数据结构核心算法实现及基本软件工程规范")

# ========== ACTIVITIES ==========
add_section_title(doc, "\U0001f3eb 在校经历与实践")

add_bullet(doc, "心理咨询室助教 (累计志愿服务 40+ 小时): 协助心理咨询师开展工作，日常与学生沟通交流，培养良好的倾听能力、同理心和沟通技巧")
add_bullet(doc, "校园心理剧参演: 获广州市一等奖、广东省二等奖，锻炼团队协作和临场应变能力")
add_bullet(doc, "自学能力: 通过官方文档、GitHub 开源社区和技术博客，独立学习 ESP32 嵌入式开发、LLM API 集成、Python 桌面应用等课外技术栈，并应用于 4 个完整项目实践")

# ========== SELF EVAL ==========
add_section_title(doc, "✨ 自我评价")

p = doc.add_paragraph()
run = p.add_run(
    "软件工程专业大二在读，对编程和产品创造有强烈热情。"
    "具备从需求分析、技术选型、架构设计到编码实现的全流程独立开发能力，"
    "已完成 4 个涵盖 AI/LLM 应用架构、嵌入式 IoT、桌面 GUI 的全栈项目。"
    "在 LLM 应用开发方向有较深入实践，能够设计多 Provider 抽象层并处理不同 API 协议间的适配转换。"
    "善于通过阅读官方文档和开源社区资源快速掌握新技术，具有较强的自驱力和问题解决能力。"
    "在团队项目中多次担任组长，具备良好的沟通协调和项目推进能力。"
)
run.font.size = Pt(10)

# Save
output_path = r'C:\Users\24522\Desktop\个人\简历\陈梓烨-个人简历-完善版.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print("DONE: " + output_path)
