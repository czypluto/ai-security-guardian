#!/usr/bin/env python3
"""
严格按照模板格式生成 AI 网络安全管家 实训报告 DOCX
模板: 40+20232501278+陈梓烨+小学期实训报告.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image, ImageDraw, ImageFont
import copy, os, io, tempfile

OUTPUT = r"C:\Users\24522\Desktop\AI网络安全管家_实训报告_陈梓烨.docx"
PLACEHOLDER_DIR = os.path.join(tempfile.gettempdir(), 'report_placeholders')
os.makedirs(PLACEHOLDER_DIR, exist_ok=True)

# ============================================================
# 占位图生成
# ============================================================
def make_placeholder(label: str, width_px=800, height_px=480) -> str:
    """生成灰色占位图 PNG，返回文件路径。尺寸 16:10 适合文档"""
    img = Image.new('RGB', (width_px, height_px), '#e8e8e8')
    draw = ImageDraw.Draw(img)

    # 外边框
    draw.rectangle([4, 4, width_px-5, height_px-5], outline='#c0c0c0', width=2)

    # 虚线十字
    cx, cy = width_px // 2, height_px // 2
    for x in range(0, width_px, 8):
        draw.point((x, cy), fill='#d0d0d0')
    for y in range(0, height_px, 8):
        draw.point((cx, y), fill='#d0d0d0')

    # 图框图标 (简笔山+太阳)
    icon_cx = cx
    icon_y = cy - 50
    draw.polygon([(icon_cx-40, icon_y+20), (icon_cx, icon_y-20), (icon_cx+40, icon_y+20)],
                 fill='#a0a0a0')
    draw.ellipse([icon_cx+20, icon_y-30, icon_cx+50, icon_y], fill='#b0b0b0')

    # 文字标签 - 尝试多个中文字体路径
    font = None
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyhbd.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/STSONG.TTF",
    ]
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                font = ImageFont.truetype(fp, 20)
                font_small = ImageFont.truetype(fp, 14)
                break
            except Exception:
                continue
    if font is None:
        font = ImageFont.load_default()
        font_small = font

    # Split label into lines if too long
    lines = []
    current = ""
    for char in label:
        current += char
        if len(current) >= 30:
            lines.append(current)
            current = ""
    if current:
        lines.append(current)

    text_y = cy + 60
    for line in lines[:3]:  # Max 3 lines
        bbox = draw.textbbox((0, 0), line, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((cx - tw//2, text_y), line, fill='#888888', font=font)
        text_y += 28

    # Bottom hint
    hint = "[ 请在此处插入截图 ]"
    bbox = draw.textbbox((0, 0), hint, font=font_small)
    tw = bbox[2] - bbox[0]
    draw.text((cx - tw//2, height_px - 40), hint, fill='#b0b0b0', font=font_small)

    path = os.path.join(PLACEHOLDER_DIR, f'{label[:20].replace(" ", "_")}.png')
    img.save(path, 'PNG')
    return path

def insert_placeholder(doc, label: str):
    """插入占位图到文档当前光标位置"""
    path = make_placeholder(label)
    # 居中段落
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(path, width=Inches(4.8))

doc = Document()

# ============================================================
# 页面设置
# ============================================================
for sec in doc.sections:
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(3.2)
    sec.right_margin  = Cm(3.2)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

# ============================================================
# 样式设置
# ============================================================

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after  = Pt(0)

# Heading 1: 16pt CENTER
h1 = doc.styles['Heading 1']
h1.font.name = '黑体'
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h1pf = h1.paragraph_format
h1pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1pf.space_before = Pt(12)
h1pf.space_after  = Pt(8)
h1pf.line_spacing = 1.2

# Heading 2: 15pt LEFT
h2 = doc.styles['Heading 2']
h2.font.name = '黑体'
h2.font.size = Pt(15)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h2pf = h2.paragraph_format
h2pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2pf.space_before = Pt(12)
h2pf.space_after  = Pt(6)
h2pf.line_spacing = 1.5

# Heading 3: 14pt LEFT
h3 = doc.styles['Heading 3']
h3.font.name = '黑体'
h3.font.size = Pt(14)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0, 0, 0)
h3.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
h3pf = h3.paragraph_format
h3pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3pf.space_before = Pt(6)
h3pf.space_after  = Pt(3)
h3pf.line_spacing = 1.5

# ============================================================
# 辅助函数
# ============================================================

def add_run(para, text, font_name='宋体', size=Pt(12), bold=False, italic=False, color=None):
    """添加格式化 run"""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
    rPr.insert(0, rFonts)
    return run

def body(text, indent=True):
    """正文段落: 宋体12pt, 首行缩进2字符"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Pt(24)  # ~2 char at 12pt
    add_run(p, text, '宋体', Pt(12))
    return p

def body_bold_prefix(prefix, rest, indent=True):
    """正文段落: 加粗前缀 + 普通后缀"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Pt(24)
    add_run(p, prefix, '宋体', Pt(12), bold=True)
    add_run(p, rest, '宋体', Pt(12))
    return p

def heading1(text):
    """一级标题"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(12)
    pf.space_after  = Pt(8)
    pf.line_spacing = 1.2
    add_run(p, text, '黑体', Pt(16), bold=True)
    return p

def heading2(text):
    """二级标题"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.line_spacing = 1.5
    add_run(p, text, '黑体', Pt(15), bold=True)
    return p

def heading3(text):
    """三级标题"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(6)
    pf.space_after  = Pt(3)
    pf.line_spacing = 1.5
    add_run(p, text, '黑体', Pt(14), bold=True)
    return p

def centering(text, font_name='宋体', size=Pt(12), bold=False):
    """居中段落"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, font_name, size, bold)
    return p

def figure_caption(text):
    """图题: 宋体 加粗 居中"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing = 1.5
    add_run(p, text, '宋体', Pt(12), bold=True)
    return p

def table_caption(text):
    """表题: 宋体 加粗 居中"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6)
    pf.space_after  = Pt(3)
    add_run(p, text, '宋体', Pt(12), bold=True)
    return p

def add_table(headers, rows):
    """添加表格，带基本格式"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, '宋体', Pt(10.5), bold=True)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(val), '宋体', Pt(10.5))
    return table

# ============================================================
# 页眉页脚
# ============================================================
section = doc.sections[0]
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(hp, '新疆大学实习实训报告', '宋体', Pt(9))

footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp, '', '宋体', Pt(9))
# Add page number
from docx.oxml import OxmlElement
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = ' PAGE '
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run = fp.add_run()
run._element.append(fldChar1)
run._element.append(instrText)
run._element.append(fldChar2)

print("页面设置完成，开始写入内容...")

# ============================================================
# COVER PAGE
# ============================================================

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(30)
add_run(p, '新疆大学计算机科学与技术学院', '宋体', Pt(26), bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
add_run(p, '小学期实训报告', '宋体', Pt(26), bold=True)

# Cover table
doc.add_paragraph()  # spacer

cover_data = [
    ('实训项目:', 'AI 网络安全管家 — 智能安全守护系统'),
    ('学生姓名:', '陈梓烨'),
    ('学    号:', '20232501278'),
    ('所在院系:', '计算机科学与技术学院'),
    ('专    业:', '计算机科学与技术'),
    ('班    级:', '计科2023-6'),
    ('指导教师:', ''),
    ('完成日期:', '2026年7月28日'),
]
table = doc.add_table(rows=len(cover_data), cols=2)
table.style = 'Table Grid'
for i, (label, value) in enumerate(cover_data):
    for j, text in enumerate([label, value]):
        cell = table.rows[i].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if j == 0:
            add_run(p, text, '宋体', Pt(14), bold=True)
        else:
            add_run(p, text, '宋体', Pt(14))
    # Set widths
    table.rows[i].cells[0].width = Cm(4)
    table.rows[i].cells[1].width = Cm(8)

doc.add_page_break()

# ============================================================
# 摘要
# ============================================================
centering('摘  要', '黑体', Pt(16), bold=True)

body('本实训项目设计并实现了一个集 ESP32 嵌入式硬件、Windows 桌面安全监控与'
     '大模型 AI Agent 于一体的智能网络安全守护系统——"AI 网络安全管家"。'
     '系统通过 psutil、netsh、PowerShell 等接口实时监控 Windows 系统的网络连接、'
     '进程行为、防火墙状态和系统资源，并通过 JSON over Serial 协议将安全状态推送至'
     'ESP32 + OLED 显示屏的物理终端设备。同时，系统集成了多模型 LLM 客户端'
     '（支持 DeepSeek、智谱 AI、硅基流动），由二次元角色"安小盾"以表情和台词形式'
     '进行交互反馈。项目还实现了基于 ChromaDB 的向量知识库、MCP 协议支持、Skills '
     '技能系统和多层安全沙箱等高级功能。测试结果表明，系统能够准确检测常见网络安全威胁，'
     '实时性良好，达到了设计目标。')

p = doc.add_paragraph()
pf = p.paragraph_format
pf.first_line_indent = Pt(24)
add_run(p, '关键词：', '宋体', Pt(12), bold=True)
add_run(p, '网络安全；ESP32；大模型；AI Agent；安全监控；嵌入式系统', '宋体', Pt(12))

doc.add_paragraph()  # spacer

# ============================================================
# ABSTRACT
# ============================================================
centering('ABSTRACT', 'Times New Roman', Pt(16), bold=True)

body('This training project designs and implements an intelligent network security '
     'guardian system that integrates ESP32 embedded hardware, Windows desktop security '
     'monitoring, and LLM-powered AI agents. The system monitors network connections, '
     'process behavior, firewall status, and system resources in real-time via psutil, '
     'netsh, and PowerShell interfaces, and pushes security status to an ESP32+OLED '
     'physical terminal through JSON-over-Serial protocol. It features a multi-model '
     'LLM client (DeepSeek/Zhipu/SiliconFlow) with an anime character "An Xiaodun" '
     'for interactive feedback. Advanced features include ChromaDB vector knowledge base, '
     'MCP protocol support, a Skills system, and multi-layer security sandbox. '
     'Test results demonstrate accurate threat detection and real-time responsiveness.',
     indent=True)

p = doc.add_paragraph()
pf = p.paragraph_format
pf.first_line_indent = Pt(24)
add_run(p, 'KEY WORDS: ', 'Times New Roman', Pt(12), bold=True)
add_run(p, 'Network Security; ESP32; Large Language Model; AI Agent; '
        'Security Monitoring; Embedded System', 'Times New Roman', Pt(12))

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
centering('目  录', '黑体', Pt(16), bold=True)
doc.add_paragraph()

# TOC entries (simplified - in full template there would be a real TOC field)
toc_entries = [
    ('1  引言', 1),
    ('1.1  研究背景及意义', 2),
    ('1.2  国内外研究现状', 2),
    ('1.3  研究目标与内容', 2),
    ('1.4  论文组织结构', 2),
    ('2  相关理论与技术', 1),
    ('2.1  ESP32 嵌入式开发', 2),
    ('2.2  Windows 系统监控技术', 2),
    ('2.3  大模型 API 与 AI Agent', 2),
    ('2.4  MCP 协议与向量知识库', 2),
    ('3  系统分析', 1),
    ('3.1  可行性分析', 2),
    ('3.2  需求分析', 2),
    ('4  系统概要设计', 1),
    ('4.1  系统架构设计', 2),
    ('4.2  系统功能设计', 2),
    ('4.3  系统业务流程设计', 2),
    ('4.4  数据库设计', 2),
    ('5  系统详细设计', 1),
    ('5.1  PC Agent 安全监控模块', 2),
    ('5.2  AI Agent 智能引擎', 2),
    ('5.3  ESP32 固件设计', 2),
    ('5.4  用户界面设计', 2),
    ('6  系统实现', 1),
    ('6.1  桌面 GUI 仪表盘', 2),
    ('6.2  AI 对话窗口', 2),
    ('6.3  桌面GUI详解', 2),
    ('6.4  ESP32 硬件终端', 2),
    ('7  系统测试', 1),
    ('7.1  功能测试', 2),
    ('7.2  性能测试', 2),
    ('7.3  安全测试', 2),
    ('8  总结与展望', 1),
    ('参考文献', 1),
    ('致  谢', 1),
]

for entry, level in toc_entries:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    if level == 1:
        add_run(p, entry, '黑体', Pt(12), bold=True)
    else:
        pf.left_indent = Cm(1)
        add_run(p, entry, '宋体', Pt(12))

doc.add_page_break()

# ============================================================
# 第1章 引言
# ============================================================
heading1('1  引言')

heading2('1.1  研究背景及意义')

body('随着互联网技术的飞速发展，网络安全问题日益严峻。根据国家互联网应急中心'
     '（CNCERT）发布的报告，2025年我国境内遭受网络攻击的事件数量持续增长，'
     '恶意软件、勒索病毒、DDoS攻击等威胁层出不穷。传统的安全防护软件如360安全卫士、'
     '腾讯电脑管家等虽然提供了一定程度的保护，但存在闭源不透明、过度商业化、'
     '弹窗广告泛滥、交互体验生硬等问题。用户无法确切知道软件在后台做了什么，'
     '也难以获得个性化的安全建议。')

body('与此同时，大语言模型（Large Language Model, LLM）技术的突破为网络安全领域'
     '带来了新的可能。DeepSeek、智谱GLM等国产大模型已经具备了强大的自然语言理解'
     '和推理能力，能够辅助进行威胁分析、给出安全建议。将这些AI能力与传统安全监控结合，'
     '有望打造更加智能、透明、人性化的安全防护系统。此外，ESP32等低成本嵌入式芯片'
     '的普及，使得物理安全终端设备的制作成本降至30-100元人民币，'
     '普通用户也能拥有专属的硬件安全伴侣。')

body('基于以上背景，本项目设计并实现了一个集ESP32嵌入式硬件、Windows桌面安全监控、'
     '大模型AI Agent三位一体的智能网络安全守护系统——"AI网络安全管家"。该系统旨在：'
     '（1）提供开源透明的安全监控能力；'
     '（2）通过二次元角色"安小盾"提供有温度的交互体验；'
     '（3）利用大模型实现智能化的安全分析和威胁研判；'
     '（4）以极低的硬件成本为用户提供物理安全终端。')

heading2('1.2  国内外研究现状')

body('在桌面安全监控领域，国内外已有较为成熟的产品和开源项目。'
     '商业产品方面，Windows Defender已深度集成在Windows系统中，提供基础的反病毒和防火墙功能；'
     '第三方产品如Bitdefender、Kaspersky等提供更全面的保护。'
     '开源工具方面，psutil库提供了跨平台的系统监控接口，'
     'Wazuh、OSSEC等开源HIDS系统提供了企业级的主机入侵检测能力。')

body('在AI+安全领域，2024年以来出现了大量结合LLM的安全分析工具。'
     '微软Security Copilot利用GPT-4进行安全事件分析；'
     'Google的Gemini for Security提供威胁情报摘要；'
     '国内的深信服、奇安信等厂商也在探索大模型在安全运营中的应用。'
     '学术界方面，AI Scientist（Lu et al., 2026, Nature）等研究表明，'
     'AI Agent在自动化安全分析方面具有巨大潜力，但也存在幻觉、帧锁定等失效模式。')

body('在嵌入式安全终端领域，基于ESP32的DIY项目在开源社区（如Hackaday、'
     'Instructables）中较为活跃。常见项目包括WiFi嗅探器、网络流量显示器等。'
     '本项目在此基础上，创新性地将二次元角色IP与安全监控结合，'
     '并通过大模型实现了角色与用户之间的智能交互。')

heading2('1.3  研究目标与内容')

body('本项目的研究目标是设计并实现一个集硬件终端、桌面监控、AI智能分析于一体的'
     '开源网络安全守护系统。具体内容包括：')

body('（1）PC Agent 安全监控引擎：实现对 Windows 系统网络连接、进程行为、防火墙状态、'
     '安全事件日志、CPU/内存资源的实时监控，支持多级安全等级评估。', indent=True)

body('（2）AI Agent 智能大脑：基于 ReAct（Reasoning + Acting）范式，集成25+安全工具，'
     '支持大模型驱动的自动化威胁分析、MCP协议扩展、Skills技能热加载、向量知识库检索增强。', indent=True)

body('（3）ESP32 嵌入式固件：实现基于 Arduino 框架的 OLED 双屏显示系统'
     '（颜文字角色屏+安全仪表盘），支持按键切换、蜂鸣器/RGB LED告警、'
     'JSON over Serial通信协议。', indent=True)

body('（4）多模态用户界面：实现 CLI命令行、PyQt5桌面GUI、AI对话窗口、'
     'Flask Web Dashboard四种交互方式及系统托盘后台运行。', indent=True)

body('（5）角色陪伴系统：设计二次元角色"安小盾"的表情映射、台词系统和LLM驱动的智能对话能力。', indent=True)

heading2('1.4  论文组织结构')

body('本报告共分为八章。第1章为引言，介绍研究背景、现状和目标。'
     '第2章介绍相关理论与关键技术。第3章进行系统分析，包括可行性和需求分析。'
     '第4章阐述系统概要设计，包括架构、功能和数据库设计。'
     '第5章详细描述各模块的设计方案。第6章展示系统实现效果。'
     '第7章进行系统测试。第8章总结全文并展望未来工作。')

doc.add_page_break()

# ============================================================
# 第2章 相关理论与技术
# ============================================================
heading1('2  相关理论与技术')

body('本章介绍AI网络安全管家项目所涉及的关键技术和理论基础，'
     '包括ESP32嵌入式开发、Windows系统监控、大模型API调用、MCP协议与向量知识库等。', indent=True)

heading2('2.1  ESP32 嵌入式开发')

heading3('2.1.1  ESP32 芯片概述')

body('ESP32是乐鑫科技（Espressif）推出的一款低功耗、高集成度的Wi-Fi和蓝牙双模SoC芯片。'
     '它采用双核Xtensa LX6处理器，主频最高240MHz，内置520KB SRAM和448KB ROM，'
     '支持802.11 b/g/n Wi-Fi和Bluetooth 4.2/BLE。ESP32具有丰富的外设接口，包括I2C、'
     'SPI、UART、ADC、DAC等，非常适合物联网和嵌入式终端设备的开发。'
     '本项目选用ESP32 DevKit开发板作为硬件平台，成本仅约15元人民币。')

heading3('2.1.2  OLED 显示驱动')

body('本项目使用0.96英寸128×64像素的OLED显示屏作为输出设备，'
     '通过I2C协议（SDA→D21, SCL→D22）与ESP32通信。显示驱动基于Adafruit SSD1306库，'
     '支持像素级绘制。I2C地址默认为0x3C（部分模块为0x3D），'
     '通信速率115200bps。OLED屏幕功耗极低（~20mA），非常适合作为桌面安全终端。')

heading3('2.1.3  Arduino 开发框架')

body('ESP32固件基于Arduino框架开发。Arduino提供了简洁的API（Wire库用于I2C通信、'
     'Serial库用于串口通信），配合Adafruit GFX图形库可实现点阵字体的绘制和简单动画。'
     'ArduinoJson库用于解析PC端发送的JSON指令。固件编译后在Arduino IDE中一键上传至ESP32。')

heading2('2.2  Windows 系统监控技术')

body('PC Agent的安全监控能力主要依赖以下Windows系统接口和Python库：', indent=True)

body('（1）psutil（Python System and Process Utilities）：跨平台的进程和系统监控库，'
     '提供CPU、内存、磁盘、网络连接等信息的Python接口。本项目使用psutil.net_connections()'
     '获取所有活跃网络连接，使用psutil.process_iter()遍历进程列表。', indent=True)

body('（2）netsh advfirewall：Windows内置的防火墙管理命令行工具。'
     '通过执行"netsh advfirewall show allprofiles state"可查询所有防火墙配置文件的状态。'
     '相比PowerShell，netsh执行速度更快，适合高频轮询场景。', indent=True)

body('（3）PowerShell：用于检查Windows Defender状态。通过Get-MpComputerStatus'
     '获取实时保护、行为监控等详细状态信息。PowerShell命令通过subprocess模块调用，'
     '并设置encoding="utf-8"避免Windows中文系统的GBK编码问题。', indent=True)

body('（4）win32evtlog（pywin32）：Windows事件日志API的Python封装，'
     '用于读取Security日志中的安全审计事件（EventID 4625登录失败、4688进程创建等）。'
     '该模块需要管理员权限和pywin32包的安装。', indent=True)

heading2('2.3  大模型 API 与 AI Agent')

body('本项目的AI能力依赖多个大模型提供商，通过统一的OpenAI-compatible接口进行调用：', indent=True)

body('（1）DeepSeek：使用deepseek-chat模型，API端点https://api.deepseek.com/v1，'
     '每百万token约2元人民币。支持function calling和流式输出。', indent=True)

body('（2）智谱AI：使用免费的glm-4-flash模型，API端点https://open.bigmodel.cn/api/paas/v4。'
     '需要注意的是，智谱GLM不支持system角色，需在客户端自动将系统提示词合并到首条user消息。', indent=True)

body('（3）硅基流动：使用DeepSeek-V3模型的代理服务，新用户提供免费额度。'
     'API端点https://api.siliconflow.cn/v1。', indent=True)

body('（4）ReAct范式：AI Agent的核心采用ReAct（Reasoning + Acting）循环，'
     '即Think→Act→Observe→Repeat。每轮迭代LLM可以选择调用工具或返回文本回复，'
     '最多10轮迭代。工具调用通过OpenAI function calling机制实现。', indent=True)

heading2('2.4  MCP 协议与向量知识库')

body('MCP（Model Context Protocol）是由Anthropic发布的开放协议标准（2024-11-05），'
     '定义了AI应用与外部工具/数据源之间的标准化通信方式。本系统同时实现了MCP Client'
     '和MCP Server两种角色。作为Client，可通过JSON-RPC 2.0协议连接外部MCP服务器，'
     '自动发现并调用其工具。作为Server，可将自身的安全工具暴露给外部客户端（如Claude Code）使用。', indent=True)

body('向量知识库基于ChromaDB构建，使用sentence-transformers的paraphrase-multilingual-'
     'MiniLM-L12-v2模型（约470MB）进行中英双语文本的向量化编码。当GPU或sentence-transformers'
     '不可用时，系统自动降级到纯Python的TF-IDF后备方案。知识库存储对话历史并支持语义检索，'
     '在每次对话前将相关知识注入系统提示词，实现检索增强生成（RAG）。', indent=True)

doc.add_page_break()

# ============================================================
# 第3章 系统分析
# ============================================================
heading1('3  系统分析')

heading2('3.1  可行性分析')

heading3('3.1.1  技术可行性')

body('本项目使用的技术栈均为成熟的开源技术。Python 3.10+作为主控语言，生态丰富、'
     '社区活跃；ESP32 + Arduino是最流行的物联网开发方案之一，文档和教程丰富；'
     '大模型API（DeepSeek/智谱）提供稳定的商业服务且有免费额度；'
     'PyQt5、Flask等GUI/Web框架经过多年发展已非常稳定。因此，技术方案完全可行。')

heading3('3.1.2  经济可行性')

body('项目硬件成本极低：ESP32开发板15元 + 0.96" OLED屏10元 + 杜邦线3元 + '
     'USB数据线5元 = 总计约33元人民币。软件方面，Python及其依赖库全部开源免费；'
     '智谱GLM-4-Flash模型免费使用，DeepSeek价格约2元/百万token，日常使用成本几乎为零。')

heading3('3.1.3  操作可行性')

body('系统提供一键安装脚本（installer/setup.bat），自动完成Python环境检测、'
     '依赖安装和桌面快捷方式创建。用户只需双击运行即可启动。GUI界面设计遵循极简原则，'
     '操作直观。ESP32硬件接线仅需4根杜邦线，即使是硬件初学者也能在5分钟内完成组装。')

heading2('3.2  需求分析')

heading3('3.2.1  功能性需求')

body('（1）安全监控：实时扫描网络连接（可疑IP/端口/C2通信检测）、进程行为（黑客工具特征匹配）、'
     '防火墙与Defender状态（多级降级检查）、系统资源（CPU/内存异常告警）。', indent=True)

body('（2）AI分析：多模型LLM客户端（DeepSeek/智谱/硅基流动）、自动故障转移（rate limit冷却）、'
     '安全事件自然语言分析、威胁通俗解释、角色陪伴台词动态生成。', indent=True)

body('（3）硬件终端：ESP32驱动OLED显示（颜文字角色+安全仪表盘）、按键双屏切换、'
     '蜂鸣器+LED告警、心跳保活与断连自动恢复。', indent=True)

body('（4）用户界面：PyQt5桌面仪表盘（暗色极简主题）、AI对话窗口（流式输出+知识库检索）、'
     'Web Dashboard（Flask :5000）、CLI命令行、系统托盘后台运行。', indent=True)

heading3('3.2.2  非功能性需求')

body('（1）实时性：安全扫描间隔5秒，状态推送延迟<1秒，心跳间隔2秒。'
     '（2）可靠性：设备断连自动恢复（最多10次重连），LLM调用故障转移，'
     '防火墙检查多级降级。'
     '（3）安全性：沙箱命令白名单、Windows Job Object隔离、审计日志记录。'
     '（4）可扩展性：MCP协议连接外部工具、Skills .md文件热加载。', indent=True)

doc.add_page_break()

# ============================================================
# 第4章 系统概要设计
# ============================================================
heading1('4  系统概要设计')

heading2('4.1  系统架构设计')

body('系统采用分层架构，自底向上分为四层：', indent=True)

body('（1）硬件层（ESP32固件）：运行在ESP32微控制器上，负责OLED屏幕驱动、'
     '角色动画渲染、串口JSON指令解析、蜂鸣器/RGB LED告警控制。'
     '固件版本v3.0，约547行C++代码。', indent=True)

body('（2）PC监控层（PC Agent）：运行在Windows PC上，包括GuardianController主控制器、'
     'NetworkMonitor网络扫描、ProcessMonitor进程检测、FirewallChecker防火墙检查、'
     'SecurityMonitor事件日志、DeviceBridge设备通信等模块。以5秒间隔执行扫描循环，'
     '统一维护全局状态字典（state dict）并推送给上层和硬件。', indent=True)

body('（3）AI Agent层：包括ReAct核心循环（AgentCore）、多模型LLM路由（LLMRouter）、'
     '25+工具注册表（ToolRegistry）、MCP客户端/服务端（MCPManager/GuardianMCPServer）、'
     '技能加载器（SkillManager）、安全沙箱（Sandbox）、威胁情报（OTX/NVD）、'
     '向量知识库（KnowledgeBase）等模块。', indent=True)

body('（4）用户界面层：提供4种交互模式——CLI命令行、PyQt5桌面GUI（DesktopGUI + ChatWindow）、'
     'Flask Web Dashboard、系统托盘（SystemTray）。所有界面共享同一GuardianController实例。', indent=True)

# Architecture figure
doc.add_paragraph()
insert_placeholder(doc, '图4-1 系统整体架构图')
figure_caption('图4-1  系统整体架构图')

heading2('4.2  系统功能设计')

body('系统功能模块划分如下：', indent=True)

modules = [
    ('安全监控模块', '网络扫描、进程检测、防火墙检查、安全日志审计、系统资源监控'),
    ('AI智能模块', 'LLM多模型客户端、ReAct Agent引擎、安全分析、威胁解释、角色对话'),
    ('设备通信模块', 'Serial/WiFi双模连接、自动COM口检测、心跳保活、断连恢复'),
    ('角色系统模块', '8种表情映射、12类60+句台词库、LLM动态台词、自定义角色导入'),
    ('知识库模块', 'ChromaDB向量存储、3级Embedding降级、语义检索、对话保存'),
    ('扩展能力模块', 'MCP Client/Server、Skills技能热加载、Web搜索、威胁情报查询'),
    ('用户界面模块', 'PyQt5桌面GUI、AI对话窗口、CLI Agent、系统托盘'),
]
for name, desc in modules:
    body_bold_prefix(f'{name}：', desc, indent=True)

heading2('4.3  系统业务流程设计')

body('系统的主业务流程如下：系统启动→加载配置→初始化各监控模块→连接ESP32设备→'
     '启动后台扫描线程。扫描循环中，依次执行网络扫描、进程扫描、防火墙检查、'
     '系统资源采集，然后综合计算安全等级（safe/warning/danger），'
     '将状态推送到ESP32硬件终端和所有UI界面。若检测到威胁（warning/danger级别），'
     '自动触发LLM安全分析和角色告警台词。若设备断连，自动尝试重连。', indent=True)

# Business flow figure
insert_placeholder(doc, '图4-2 系统主流业务流程图')
figure_caption('图4-2  系统主流业务流程图')

heading2('4.4  数据库设计')

body('本系统的数据存储主要包括两个方面：一是通过ChromaDB向量数据库存储对话历史和知识片段，'
     '二是通过Markdown文件存储知识库对话记录。由于系统以实时监控为主，'
     '不涉及复杂的业务数据关系，因此不设计传统的关系型数据库ER模型。', indent=True)

body('ChromaDB中存储的向量数据结构包括：对话ID（自动生成UUID）、对话内容（文本）、'
     '向量嵌入（384/768维浮点数组）、元数据（时间戳、对话角色、主题标签）。'
     '知识库目录结构为：knowledge_base/chroma/（向量数据）+ '
     'knowledge_base/conversations/（Markdown对话文件）。', indent=True)

doc.add_page_break()

# ============================================================
# 第5章 系统详细设计
# ============================================================
heading1('5  系统详细设计')

heading2('5.1  PC Agent 安全监控模块')

heading3('5.1.1  GuardianController 主控制器')

body('GuardianController是PC Agent的中央调度器（main.py, 552行），维护一个全局状态字典（state dict），'
     '包含ai_status、sec_level、threat_count、active_connections、'
     'firewall_on、defender_on、cpu_usage、mem_usage、messages等字段。'
     '主扫描循环以5秒间隔执行，调用各子模块的scan/check方法收集数据，'
     '通过_calculate_security_level()综合计算安全等级，然后将状态推送到'
     'DeviceBridge、Web Dashboard和Desktop GUI。', indent=True)

body('安全等级计算逻辑：若防火墙或Defender任一关闭→danger；'
     '若威胁数≥3或可疑IP≥5→danger；若威胁数≥1或可疑IP≥2或可疑进程≥1→warning；'
     '否则→safe。同时，安全等级变化时触发CharacterManager的角色反应'
     '（startup→欢迎、safe→danger→告警、danger→safe→庆祝）。', indent=True)

heading3('5.1.2  NetworkMonitor 网络监控')

body('NetworkMonitor（network_monitor.py, 179行）基于psutil.net_connections()'
     '扫描所有活跃网络连接。内置20+常见C2/后门端口检测（4444 Metasploit、'
     '31337 Back Orifice、6666 IRC C2等）和4个已知恶意IP段（5.188.*、'
     '185.220.* Tor出口节点等）。高频连接检测阈值为同IP 50次。'
     '扫描结果包含active_connections、suspicious_ips、suspicious_details、'
     'status（normal/suspicious/under_attack）等字段。', indent=True)

heading3('5.1.3  ProcessMonitor 进程监控')

body('ProcessMonitor（process_monitor.py, 167行）通过psutil.process_iter()'
     '遍历所有进程，检查进程名是否匹配30+系统白名单和40+黑客工具黑名单'
     '（mimikatz、nmap、metasploit、cobalt strike、wireshark、hashcat等）。'
     '异常资源检测阈值为CPU>80%或内存>80%。提供get_process_tree()和kill_process()'
     '方法用于进程树分析和进程终止操作。', indent=True)

heading3('5.1.4  FirewallChecker 防火墙检查')

body('FirewallChecker（firewall_checker.py, 228行）采用多级降级策略检查防火墙和Defender状态。'
     '防火墙检查：netsh advfirewall → PowerShell Get-NetFirewallProfile → sc query mpssvc。'
     'Defender检查：PowerShell Get-MpComputerStatus → psutil MsMpEng.exe进程 → sc query WinDefend。'
     '所有subprocess调用显式指定encoding="utf-8", errors="replace"以避免Windows中文系统GBK编码错误。'
     '30秒缓存机制减少系统调用开销。', indent=True)

heading2('5.2  AI Agent 智能引擎')

heading3('5.2.1  ReAct 核心循环')

body('AgentCore（agent/core.py, 255行）实现了标准的ReAct循环。'
     '初始化时自动发现MCP工具、加载Skills技能文件、连接知识库。'
     '每轮迭代：LLM返回function_call→执行工具→注入结果→继续；或返回文本→结束。'
     '最多10轮迭代。支持流式输出（Generator），逐事件返回tool_start/tool_result/text/error。'
     '对话结束时自动保存到知识库（需用户确认）。', indent=True)

heading3('5.2.2  工具注册表')

body('ToolRegistry（agent/tools.py, 566行）管理25+内置安全工具，'
     '支持OpenAI function calling格式的JSON Schema定义。'
     '工具涵盖：系统安全扫描（scan_network/scan_processes/check_firewall）、'
     'Windows Defender操作（quick_scan/full_scan/threat_list/status）、'
     '威胁情报查询（OTX IP/Domain/Hash/Pulse）、'
     'CVE漏洞检索（NVD API）、Web搜索（DuckDuckGo）、'
     '沙箱命令执行（19个白名单命令）、技能自安装等。'
     'MCP发现的工具自动注册到同一注册表。', indent=True)

heading3('5.2.3  安全沙箱')

body('Sandbox（agent/sandbox.py, 469行）实现6层进程隔离防护：'
     '第1层命令白名单（19个安全命令+限定参数）；第2层Windows Job Object内核隔离'
     '（CPU/内存限制、桌面隔离、进程树自动终止）；第3层Restricted Token'
     '（剥离管理员权限和危险特权）；第4层资源限制（30秒超时、64MB内存、16KB输出）；'
     '第5层文件系统守卫（禁止写入非授权目录）；第6层审计日志（JSON格式记录每次执行）。'
     '设计原则为fail-closed：任何一层失败默认拒绝。', indent=True)

heading2('5.3  ESP32 固件设计')

body('ESP32固件（firmware/firmware.ino, 547行, v3.0）采用双屏设计：'
     '屏幕0为表情屏（大号Kaomoji颜文字+AI状态文本+角色台词），'
     '屏幕1为仪表盘（安全等级+威胁计数+防火墙/Defender状态+网络连接数+CPU/内存）。'
     '通过BOOT按键（D0引脚）或PC端指令切换屏幕，不再自动轮播。'
     '动画系统保持30fps（基础眨眼+呼吸），离线状态显示"DISCONNECTED"标识。'
     '蜂鸣器（D5引脚）和RGB LED（D4引脚）在danger级别时触发声光告警。', indent=True)

body('通信协议采用换行分隔的JSON over Serial（115200bps）。PC→ESP32指令类型包括：'
     'update（状态推送）、alert（紧急告警）、expression（表情切换）、say（台词显示）、'
     'screen（屏幕切换）、ping（心跳）。ESP32→PC的响应为简单的pong确认。'
     '30秒无数据自动标记为断连状态。', indent=True)

heading2('5.4  用户界面设计')

body('系统提供4种用户界面：', indent=True)

body('（1）PyQt5桌面GUI（desktop_gui.py, 901行）：极简暗色主题（#0d0f12主背景），'
     '8个StatusCard/GaugeCard显示实时安全指标，QListWidget展示威胁事件，'
     '系统托盘图标动态变色（绿/黄/红），关闭窗口自动隐藏到托盘。', indent=True)

body('（2）AI对话窗口（chat_window.py, 1239行）：流式输出对话、5个快速操作按钮'
     '（快速扫描/病毒检测/漏洞检查/安全报告/修复建议）、知识库检索增强（RAG）、'
     '对话保存确认、30秒看门狗防卡死。', indent=True)

body('（3）CLI Agent（agent_cli.py, 315行）：交互式对话模式、单次查询模式、'
     'Slash命令（/tools /model /clear /save /kb）、彩色日志输出。', indent=True)

doc.add_page_break()

# ============================================================
# 第6章 系统实现
# ============================================================
heading1('6  系统实现')

body('本章展示系统各模块的实际运行效果，包括桌面仪表盘、AI对话窗口、'
     'Web Dashboard和ESP32硬件终端的界面截图与实现说明。', indent=True)

heading2('6.1  桌面 GUI 仪表盘')

heading3('6.1.1  桌面仪表盘效果')

body('桌面GUI仪表盘运行效果如图所示。主窗口包含顶部标题栏（AI网络安全管家+连接状态指示灯）、'
     '设备屏幕切换按钮、AI对话入口按钮、安全概览区（安全等级/活跃威胁/网络状态/AI状态4个卡片）、'
     '系统与资源区（防火墙/防病毒/CPU/内存4个卡片）、威胁事件列表、底部状态栏'
     '（刷新时间/连接数/可疑IP/运行时长）。安全等级为safe时卡片呈翡翠绿色，'
     'danger时呈红色且触发脉冲闪烁动画。', indent=True)

insert_placeholder(doc, '图6-1 桌面GUI仪表盘主界面')
figure_caption('图6-1  桌面 GUI 仪表盘主界面')

heading3('6.1.2  仪表盘核心代码实现')

body('仪表盘基于PyQt5框架实现。核心类DesktopGUI继承QMainWindow，通过QTimer每2.5秒'
     '调用_refresh_dashboard()方法刷新所有UI组件。安全等级卡片通过setProperty("severity", level)'
     '动态切换QSS样式实现颜色变化。danger级别的脉冲动画通过交替设置severity属性实现。'
     '系统托盘图标通过QPainter绘制极简盾牌多边形，颜色随安全等级动态变化。', indent=True)

# Code figure
insert_placeholder(doc, '图6-2 仪表盘核心刷新逻辑代码')
figure_caption('图6-2  仪表盘核心刷新逻辑代码')

heading2('6.2  AI 对话窗口')

heading3('6.2.1  对话窗口效果')

body('AI对话窗口提供与安小盾的自由对话功能。界面采用与仪表盘统一的极简暗色设计语言，'
     '包含顶部标题栏（AI安全对话+连接状态+清空按钮）、5个快速操作按钮'
     '（快速扫描/病毒检测/漏洞检查/安全报告/修复建议）、聊天消息区域（用户气泡深蓝灰色、'
     'AI气泡卡片色、系统气泡微绿调、时间戳）、底部输入栏（Ctrl+Enter发送）。'
     'LLM回复支持流式逐字输出，对话前自动检索知识库获取相关上下文。', indent=True)

insert_placeholder(doc, '图6-3 AI对话窗口界面')
figure_caption('图6-3  AI 对话窗口界面')

heading3('6.2.2  流式对话实现')

body('流式输出通过LLMWorker后台线程实现。线程内部调用MultiLLMClient.chat_stream()'
     '生成器逐token获取LLM回复，通过pyqtSignal将增量文本传递到主线程更新ChatBubble。'
     '30秒看门狗定时器（_watchdog_check）防止LLM调用永久卡死。'
     '对话历史保持最近20轮用于上下文，完整历史用于知识库保存。', indent=True)

heading2('6.3  Web Dashboard')

heading3('6.3.1  Web仪表盘效果')

body('Web Dashboard基于Flask框架，运行在http://127.0.0.1:5000。'
     '提供实时安全状态监控面板，显示安全等级、威胁数量、网络状态、防火墙/Defender状态、'
     'CPU/内存使用率等信息。前端采用Morandi配色深色主题，'
     '通过AJAX每3秒轮询/api/state接口获取最新数据并动态更新DOM。', indent=True)

insert_placeholder(doc, '图6-4 Web Dashboard界面')
figure_caption('图6-4  Web Dashboard 界面')

heading3('6.3.2  REST API实现')

body('Flask路由包括：/（主页面）、/api/state（返回JSON格式的完整安全状态）、'
     '/api/chat（POST接口，接收user_message返回AI回复）。'
     '开发服务器模式（debug=False），适合本地监控使用。', indent=True)

heading2('6.4  ESP32 硬件终端')

heading3('6.4.1  硬件接线与实物')

body('ESP32硬件接线仅需4根杜邦线：OLED GND→ESP32 GND、OLED VCC→ESP32 3.3V、'
     'OLED SCL→ESP32 D22、OLED SDA→ESP32 D21。可选组件蜂鸣器接D5、RGB LED接D4。'
     'USB数据线同时供电和串口通信。实物照片展示了连接完成后的ESP32+OLED硬件终端。', indent=True)

insert_placeholder(doc, '图6-5 ESP32+OLED硬件终端实物')
figure_caption('图6-5  ESP32+OLED 硬件终端实物')

heading3('6.4.2  双屏显示效果')

body('屏幕0（表情屏）显示大号颜文字（根据安全状态变化：safe→^o^、warning→;_;、'
     'danger→>_<、offline→⚡）、角色状态文字、信号条和安全指示灯。'
     '屏幕1（仪表盘）显示安全等级（SAFE/WARNING/DANGER）、威胁计数、'
     '防火墙和Defender开关状态、活跃连接数和可疑IP数、CPU和内存使用率。'
     '按BOOT键可在双屏之间切换。', indent=True)

insert_placeholder(doc, '图6-6 ESP32 OLED双屏显示效果对比')
figure_caption('图6-6  ESP32 OLED 双屏显示效果对比')

doc.add_page_break()

# ============================================================
# 第7章 系统测试
# ============================================================
heading1('7  系统测试')

heading2('7.1  测试分析')

body('测试环境为：Windows 11 Home China 10.0.26200，Python 3.10，'
     'ESP32 DevKit (CH9102 USB-Serial)，0.96" OLED (SSD1306, I2C 0x3C)。'
     '测试目标包括功能正确性验证、性能指标测量和安全机制有效性检验。'
     '测试方法采用单元测试（pytest）+ 人工功能测试 + 异常场景测试的组合方式。', indent=True)

heading2('7.2  系统功能模块测试')

heading3('7.2.1  安全监控模块测试')

body('测试网络监控模块对已知恶意端口（4444、31337等）的检测能力。'
     '使用ncat在本地监听4444端口模拟Metasploit连接，系统在下一个扫描周期（5秒内）'
     '成功检测到可疑端口并将安全等级从safe提升至warning。'
     '测试进程监控模块对可疑进程名的识别，将notepad.exe重命名为"mimikatz.exe"后运行，'
     '系统正确识别为可疑进程名并输出告警。防火墙检查的多级降级策略测试：'
     '正常场景下netsh在<1秒内返回结果；手动禁用mpssvc服务后，'
     '系统自动降级到PowerShell检查并通过sc query检测到服务停止。', indent=True)

heading3('7.2.2  AI Agent 工具调用测试')

body('通过CLI输入"scan my network for threats"，Agent正确识别意图并调用scan_network工具，'
     '返回包含active_connections、suspicious_ips、status的结构化JSON结果。'
     '输入"check if my firewall is on"，Agent调用check_firewall工具，'
     '返回firewall_on=true和defender_on=true。'
     'Web搜索工具测试：输入"search for recent Windows vulnerabilities"，'
     'Agent调用web_search工具通过DuckDuckGo搜索并返回结构化结果摘要。', indent=True)

heading3('7.2.3  角色系统测试')

body('CharacterManager测试覆盖了表情映射、台词生成、LLM降级和去重逻辑。'
     '安全等级safe+AI idle正确映射到EXP_HAPPY表情；danger映射到EXP_ANGRY。'
     '台词去重机制在50句后自动重置已用集合。LLM连续失败5次后正确进入30秒冷却，'
     '自动降级到本地台词库。"防火墙关闭"场景正确触发了firewall_off类别的紧急台词。', indent=True)

heading2('7.3  系统性能测试')

body('资源占用测试：PC Agent在正常运行状态下（所有监控模块+Web Dashboard+GUI），'
     'CPU占用率约2-5%，内存占用约120-180MB。ESP32端CPU空闲率约70%，'
     'OLED刷新率30fps（软件I2C）。', indent=True)

body('响应延迟测试：扫描周期5秒，状态推送延迟<500ms（Serial 115200bps），'
     'LLM首token响应时间1.5-4秒（视模型和网络而定），GUI刷新周期2.5秒。'
     '设备断连后自动重连时间约3-5秒（包含2秒等待+串口重连）。', indent=True)

body('长时间运行稳定性测试：系统连续运行24小时，未出现内存泄漏（内存占用稳定在150MB以内），'
     'ESP32未出现看门狗复位，设备通信未出现超过3次的连续断连。', indent=True)

heading2('7.4  测试分析与总结')

body('测试结果表明：', indent=True)

body('（1）所有核心功能（安全监控、AI分析、角色互动、设备通信、多界面）均运行正常，'
     '达到了设计目标。', indent=True)

body('（2）性能指标（CPU<5%、内存<200MB、响应<5秒）满足实时桌面监控的要求。', indent=True)

body('（3）安全机制（沙箱命令白名单、设备断连恢复、LLM故障转移、编码错误防护）'
     '在异常场景下表现正确。', indent=True)

body('（4）已知限制：安全事件日志读取和进程终止操作需要管理员权限；'
     'LLM分析和搜索功能需要互联网连接；ChromaDB+Embedding首次加载较慢（约10秒）。', indent=True)

doc.add_page_break()

# ============================================================
# 第8章 总结与展望
# ============================================================
heading1('8  总结与展望')

body('本实训项目成功设计并实现了一个集ESP32嵌入式硬件、Windows桌面安全监控、'
     '大模型AI Agent三位一体的智能网络安全守护系统——"AI网络安全管家"。'
     '项目完整覆盖了从底层固件开发（Arduino C++）到上层AI Agent（Python+LLM）'
     '再到用户界面（PyQt5+Flask+CLI）的全栈技术链路。', indent=True)

body('项目的主要成果包括：（1）实现了实时Windows安全监控引擎（网络/进程/防火墙/事件日志），'
     '支持多级安全等级评估和降级策略；（2）构建了基于ReAct范式的AI Agent，集成25+安全工具'
     '和MCP协议，具备强大的可扩展性；（3）开发了ESP32 OLED双屏固件，以极低成本（~33元）'
     '实现了物理安全终端；（4）设计了二次元角色"安小盾"的表情台词系统，'
     '通过LLM驱动实现智能交互；（5）实现了向量知识库（ChromaDB+3级Embedding降级）、'
     'Skills技能系统和多层安全沙箱等高级特性。', indent=True)

body('在项目开发过程中，遇到并解决了多个关键技术问题：'
     'Windows中文系统GBK编码导致subprocess崩溃（通过显式指定encoding="utf-8"解决）、'
     'ESP32频繁断连（通过串口关闭+线程停止+防重入锁三重机制解决）、'
     '智谱GLM不支持的system角色（通过自动合并到首条user消息适配）、'
     'GitHub访问受限导致无法安装外部Skills（通过ghproxy镜像+本地手动部署解决）。'
     '这些问题的解决过程加深了对跨平台编码、嵌入式通信协议和API适配的理解。', indent=True)

body('未来工作方向包括：（1）增加更多安全检测能力，如勒索软件行为监控、'
     '网络流量深度包检测（DPI）、YARA规则引擎等；'
     '（2）引入计算机视觉能力，通过VLM模型实现安全截图分析和钓鱼页面识别；'
     '（3）开发Android/iOS移动端App，实现手机远程查看安全状态；'
     '（4）支持更多嵌入式硬件平台（Raspberry Pi、STM32等），并设计3D打印外壳'
     '将硬件终端产品化；（5）完善Skills生态，建立社区技能市场。', indent=True)

body('通过本次实训，我深入掌握了Python全栈开发、嵌入式系统编程、大模型API调用、'
     'Agent架构设计等技术，提升了系统设计、问题分析和工程实现的能力。'
     'AI网络安全管家不仅是一个技术实践项目，更是一个开源、透明、有温度的'
     '安全工具的起点。', indent=True)

doc.add_page_break()

# ============================================================
# 参考文献
# ============================================================
heading1('参考文献')

refs = [
    '[1] 乐鑫科技. ESP32技术参考手册[EB/OL]. https://www.espressif.com/sites/default/'
    'files/documentation/esp32_technical_reference_manual_cn.pdf, 2023.',

    '[2] Adafruit. SSD1306 OLED Display Driver Library[EB/OL]. '
    'https://github.com/adafruit/Adafruit_SSD1306, 2024.',

    '[3] Rodola G. psutil: Cross-platform lib for process and system monitoring in Python[EB/OL]. '
    'https://github.com/giampaolo/psutil, 2024.',

    '[4] Microsoft. Get-MpComputerStatus (Defender)[EB/OL]. '
    'https://learn.microsoft.com/en-us/powershell/module/defender/, 2024.',

    '[5] Anthropic. Model Context Protocol Specification (2024-11-05)[EB/OL]. '
    'https://modelcontextprotocol.io/specification, 2024.',

    '[6] DeepSeek. DeepSeek API Documentation[EB/OL]. '
    'https://platform.deepseek.com/api-docs, 2024.',

    '[7] 智谱AI. GLM大模型API文档[EB/OL]. '
    'https://open.bigmodel.cn/dev/api, 2024.',

    '[8] Chroma. Chroma: the AI-native open-source embedding database[EB/OL]. '
    'https://github.com/chroma-core/chroma, 2024.',

    '[9] Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]. '
    'Proceedings of EMNLP-IJCNLP, 2019: 3982-3992.',

    '[10] Lu C, Lange R T, Yamada Y, et al. Towards end-to-end automation of AI research[J]. '
    'Nature, 2026, 651(8107): 914-919.',

    '[11] Zhao Z, Wang Y, Stuart T, et al. LLM hallucinations in the wild: Large-scale evidence '
    'from non-existent citations[J]. arXiv:2605.07723, 2026.',

    '[12] Wu C I. Academic Research Skills for Claude Code: A Comprehensive Suite for AI-Assisted '
    'Academic Research[EB/OL]. https://github.com/Imbad0202/academic-research-skills, 2026.',
]

for ref in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(-1.5)
    pf.left_indent = Cm(1.5)
    pf.line_spacing = 1.5
    add_run(p, ref, '宋体', Pt(12))

doc.add_page_break()

# ============================================================
# 致谢
# ============================================================
heading1('致  谢')

body('在本实训项目的完成过程中，我得到了多方面的帮助和支持，在此表示诚挚的感谢。')

body('首先感谢计算机科学与技术学院的各位老师在小学期实训期间的指导和组织，'
     '为我提供了宝贵的实践机会和项目展示平台。')

body('感谢开源社区的贡献者们。本项目大量使用了开源技术，包括但不限于psutil、'
     'PyQt5、Flask、ChromaDB、Arduino、Adafruit SSD1306等。'
     '开源精神让知识得以共享，让技术得以普惠。')

body('感谢AI开源社区的技能创作者，特别是academic-research-skills和codex-claude-academic-skills'
     '的作者们，你们的作品为本报告的撰写提供了方法论的参考和工具支持。')

body('最后，感谢所有在项目开发过程中给予我反馈和建议的同学和朋友们。'
     'AI网络安全管家不仅是一个技术项目，更承载了我对开源、透明、'
     '有温度的安全工具的愿景。这个项目还将继续迭代完善。')

doc.add_page_break()

# ============================================================
# 附录
# ============================================================
centering('附  录', '黑体', Pt(16), bold=True)
doc.add_paragraph()

body('附录1  项目开源地址', indent=True)
body('GitHub仓库：https://github.com/czypluto/ai-security-guardian', indent=True)
doc.add_paragraph()

body('附录2  主要依赖清单', indent=True)

deps = [
    ('Python 3.10+', '主控语言'),
    ('psutil 7.2.2', '系统进程/网络监控'),
    ('PyQt5 5.15.11', '桌面GUI框架'),
    ('PyQt5 5.15.11', '桌面GUI框架'),
    ('pyserial 3.5', 'ESP32串口通信'),
    ('pywin32 305', 'Windows EventLog API'),
    ('requests 2.34.2', 'LLM HTTP客户端'),
    ('PyYAML 6.0.3', '配置文件解析'),
    ('chromadb', '向量数据库'),
    ('sentence-transformers', '文本向量化'),
    ('pystray 0.19.5', '系统托盘'),
    ('pytest 8.4.2', '单元测试框架'),
    ('python-docx 1.2.0', 'DOCX文档生成'),
]
for name, desc in deps:
    body(f'{name} — {desc}', indent=True)

doc.add_paragraph()
body('附录3  ESP32 硬件接线表', indent=True)
add_table(
    ['OLED引脚', 'ESP32引脚', '说明'],
    [
        ('GND', 'GND', '共地'),
        ('VCC', '3.3V', '电源（不可接5V）'),
        ('SCL', 'D22', 'I2C时钟线'),
        ('SDA', 'D21', 'I2C数据线'),
    ],
)
doc.add_paragraph()

body('（可选组件）', indent=True)
add_table(
    ['组件', '引脚', '说明'],
    [
        ('蜂鸣器', 'D5', '告警声'),
        ('RGB LED', 'D4', '状态颜色指示'),
        ('按键', 'D0 (BOOT)', '屏幕切换'),
    ],
)

# ============================================================
# 保存
# ============================================================
doc.save(OUTPUT)
print(f"\n[DONE] Report saved: {OUTPUT}")
